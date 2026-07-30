from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os 
import subprocess

BLACK = RGBColor(0, 0, 0)
FONT = "Calibri"


def set_font(run, size, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    run.font.bold = bold


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    add_bottom_border(p)
    run = p.add_run(text.upper())
    set_font(run, 11, bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, 10)
    return p


def build_doc(cv, output_path="outputs/CV_output.docx"):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.3)
        section.right_margin = Cm(1.3)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)

    # NAME
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(cv["name"])
    set_font(run, 18, bold=True)
    run.font.all_caps = False
    run.font.small_caps = False

    # TITLE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(cv["title"])
    set_font(run, 11)

    # CONTACT LINE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    contact_parts = [cv["location"], cv["phone"], cv["email"]]
    run = p.add_run("  |  ".join(contact_parts) + "  |  ")
    set_font(run, 9)
    add_hyperlink(p, "LinkedIn", cv["linkedin"])
    run = p.add_run("  |  ")
    set_font(run, 9)
    add_hyperlink(p, "GitHub", cv["github"])
    run = p.add_run("  |  ")
    set_font(run, 9)
    add_hyperlink(p, "Portfolio", cv["portfolio"])

    # SUMMARY
    section_header(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(cv["summary"])
    set_font(run, 10)

    # EXPERIENCE
    section_header(doc, "Experience")
    for exp in cv["experience"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(exp["title"])
        set_font(run, 10, bold=True)
        run = p.add_run(f"  |  {exp['company']}  |  {exp['dates']}  |  {exp['location']}")
        set_font(run, 10)
        for b in exp["bullets"]:
            add_bullet(doc, b)

    # PROJECTS
    section_header(doc, "Projects")
    for proj in cv["projects"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(proj["name"])
        set_font(run, 10, bold=True)
        run = p.add_run(f"  |  {proj['tools']}  |  ")
        set_font(run, 10)
        if proj.get("github"):
            add_hyperlink(p, "GitHub", proj["github"])
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(proj["description"])
        set_font(run, 10)

    # EDUCATION
    section_header(doc, "Education")
    for ed in cv["education"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(ed["degree"])
        set_font(run, 10, bold=True)
        run = p.add_run(f"  |  {ed['institution']}  |  {ed['dates']}")
        set_font(run, 10)

    # SKILLS
    section_header(doc, "Skills")
    skill_rows = [
        ("Technical Skills", cv["skills"].get("technical", "")),
        ("Tools and Platforms", cv["skills"].get("tools", "")),
        ("Programming Languages and Libraries", cv["skills"].get("languages", "")),
        ("Certifications", cv["skills"].get("certifications", ""))
    ]
    for label, value in skill_rows:
        print(f"[Template] Skill row: {label} = '{value}'")
        if value and value.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(f"{label}: ")
            set_font(run, 10, bold=True)
            run = p.add_run(value)
            set_font(run, 10)

    doc.save(output_path)
        
    # Verify it actually saved
    if not os.path.exists(output_path):
        raise Exception(f"Failed to save document to {output_path}")
    
    print(f"[Bob] Verified file exists at {output_path}")
    return output_path


def convert_to_pdf(docx_path):
    output_dir = os.path.dirname(docx_path)
    result = subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
        capture_output=True,
        text=True
    )
    pdf_path = docx_path.replace(".docx", ".pdf")
    if os.path.exists(pdf_path):
        print(f"[Template] PDF saved to {pdf_path}")
        return pdf_path
    else:
        print(f"[Template] PDF conversion failed: {result.stderr}")
        return None
    
